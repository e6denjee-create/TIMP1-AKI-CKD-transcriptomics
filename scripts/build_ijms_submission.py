"""Build the IJMS manuscript, supplements, and public-repository payload."""

from __future__ import annotations

import csv
import json
import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path(r"C:\Users\dell\Downloads\TIMP1_manuscript_submission_ready.docx")
OUT = ROOT / "submission_package" / "ijms_submission_2026-06-14"
REPO = ROOT / "public_repository" / "TIMP1_AKI_CKD_transcriptomics"
ZENODO = ROOT / "public_repository" / "TIMP1_AKI_CKD_transcriptomics_zenodo_payload"
REFERENCE_LIBRARY = (
    ROOT
    / "submission_package"
    / "ijms_submission_2026-06-13"
    / "references_verified.csv"
)

TITLE = (
    "Cross-Cohort Transcriptomic Analysis Associates TIMP1 with Tubular "
    "Injury-Repair, Extracellular Matrix Remodeling, and a Reproducible "
    "Disease-State Program in Human Kidney Disease"
)
AFFILIATION = (
    "Shanxi Bethune Hospital, Shanxi Academy of Medical Sciences, Third Hospital "
    "of Shanxi Medical University, Tongji Shanxi Hospital, Taiyuan 030032, China"
)
ADDRESS = (
    "No. 99 Longcheng Avenue, Xiaodian District, Taiyuan City, "
    "Shanxi Province 030032, China"
)
GITHUB_URL = "https://github.com/e6denjee-create/TIMP1-AKI-CKD-transcriptomics"
GITHUB_RELEASE_URL = f"{GITHUB_URL}/releases/tag/v1.1.0"
ZENODO_DOI = "10.5281/zenodo.20680931"
ZENODO_DOI_URL = f"https://doi.org/{ZENODO_DOI}"

RECENT_PMIDS = [
    "39298548",
    "36142787",
    "32571916",
    "39110788",
    "36932062",
    "33176333",
    "35491858",
    "36265491",
    "38513647",
    "35064106",
    "38977708",
    "35922662",
    "38126209",
    "35513123",
    "41286490",
    "35709763",
    "38351505",
    "37130011",
    "26768243",
    "39113797",
    "32303283",
    "33073587",
    "31000567",
    "39016438",
    "39572154",
    "38609039",
    "34560077",
    "38086793",
    "36229672",
    "33411069",
]

CLASSIC_REFERENCES = [
    "Venkatachalam, M.A.; Weinberg, J.M.; Kriz, W.; Bidani, A.K. Failed Tubule Recovery, AKI-CKD Transition, and Kidney Disease Progression. J. Am. Soc. Nephrol. 2015, 26, 1765-1776. https://doi.org/10.1681/ASN.2015010006.",
    "Ferenbach, D.A.; Bonventre, J.V. Mechanisms of Maladaptive Repair after AKI Leading to Accelerated Kidney Ageing and CKD. Nat. Rev. Nephrol. 2015, 11, 264-276. https://doi.org/10.1038/nrneph.2015.3.",
    "Bonventre, J.V.; Yang, L. Cellular Pathophysiology of Ischemic Acute Kidney Injury. J. Clin. Investig. 2011, 121, 4210-4221. https://doi.org/10.1172/JCI45161.",
    "Humphreys, B.D.; Valerius, M.T.; Kobayashi, A.; Mugford, J.W.; Soeung, S.; Duffield, J.S.; McMahon, A.P.; Bonventre, J.V. Intrinsic Epithelial Cells Repair the Kidney after Injury. Cell Stem Cell 2008, 2, 284-291. https://doi.org/10.1016/j.stem.2008.01.014.",
    "Liu, Y. Cellular and Molecular Mechanisms of Renal Fibrosis. Nat. Rev. Nephrol. 2011, 7, 684-696. https://doi.org/10.1038/nrneph.2011.149.",
    "Yang, L.; Besschetnova, T.Y.; Brooks, C.R.; Shah, J.V.; Bonventre, J.V. Epithelial Cell Cycle Arrest in G2/M Mediates Kidney Fibrosis after Injury. Nat. Med. 2010, 16, 535-543. https://doi.org/10.1038/nm.2144.",
    "Canaud, G.; Brooks, C.R.; Kishi, S.; Taguchi, K.; Nishimura, K.; Magassa, S.; Scott, A.; Hsiao, L.L.; Ichimura, T.; Terzi, F.; et al. Cyclin G1 and TASCC Regulate Kidney Epithelial Cell G2-M Arrest and Fibrotic Maladaptive Repair. Sci. Transl. Med. 2019, 11, eaav4754. https://doi.org/10.1126/scitranslmed.aav4754.",
    "Rabb, H.; Griffin, M.D.; McKay, D.B.; Swaminathan, S.; Pickkers, P.; Rosner, M.H.; Kellum, J.A.; Ronco, C. Inflammation in AKI: Current Understanding, Key Questions, and Knowledge Gaps. J. Am. Soc. Nephrol. 2016, 27, 371-379. https://doi.org/10.1681/ASN.2015030261.",
    "Brew, K.; Nagase, H. The Tissue Inhibitors of Metalloproteinases (TIMPs): An Ancient Family with Structural and Functional Diversity. Biochim. Biophys. Acta 2010, 1803, 55-71. https://doi.org/10.1016/j.bbamcr.2010.01.003.",
    "Lake, B.B.; Menon, R.; Winfree, S.; Hu, Q.; Melo Ferreira, R.; Kalhor, K.; Barwinska, D.; Otto, E.A.; Ferkowicz, M.; Diep, D.; et al. An Atlas of Healthy and Injured Cell States and Niches in the Human Kidney. Nature 2023, 619, 585-594. https://doi.org/10.1038/s41586-023-05769-3.",
    "Meng, X.M.; Nikolic-Paterson, D.J.; Lan, H.Y. TGF-beta: The Master Regulator of Fibrosis. Nat. Rev. Nephrol. 2016, 12, 325-338. https://doi.org/10.1038/nrneph.2016.48.",
]

DATASET_REFERENCES = [
    "National Center for Biotechnology Information. Gene Expression Omnibus, GSE139061: Transcriptomic Signatures of Kidney Injury in Human Renal Biopsy Specimens. Available online: https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE139061 (accessed on 14 June 2026).",
    "Famulski, K.S.; de Freitas, D.G.; Kreepala, C.; Chang, J.; Sellares, J.; Sis, B.; Einecke, G.; Mengel, M.; Reeve, J.; Halloran, P.F. Molecular Phenotypes of Acute Kidney Injury in Kidney Transplants. J. Am. Soc. Nephrol. 2012, 23, 948-958. https://doi.org/10.1681/ASN.2011090887.",
    "Nakagawa, S.; Nishihara, K.; Miyata, H.; Shinke, H.; Tomita, E.; Kajiwara, M.; Matsubara, T.; Iehara, N.; Igarashi, Y.; Yamada, H.; et al. Molecular Markers of Tubulointerstitial Fibrosis and Tubular Cell Damage in Patients with Chronic Kidney Disease. PLoS ONE 2015, 10, e0136994. https://doi.org/10.1371/journal.pone.0136994.",
    "Liu, J.; Nair, V.; Zhao, Y.Y.; Chang, D.Y.; Limonte, C.; Bansal, N.; Fermin, D.; Eichinger, F.; Tanner, E.C.; Bellovich, K.A.; et al. Multi-Scalar Data Integration Links Glomerular Angiopoietin-Tie Signaling Pathway Activation with Progression of Diabetic Kidney Disease. Diabetes 2022, 71, 2664-2676. https://doi.org/10.2337/db22-0169.",
    "Hinze, C.; Kocks, C.; Leiz, J.; Karaiskos, N.; Boltengagen, A.; Cao, S.; Himmerkus, N.; Hackermuller, J.; Mertins, P.; Haller, H.; et al. Single-cell transcriptomics reveals common epithelial response patterns in human acute kidney injury. Genome Med. 2022, 14, 103. https://doi.org/10.1186/s13073-022-01106-9.",
]


def clean_text(value: object) -> str:
    return str(value).replace("¦Ā", "beta").replace("β", "beta").strip()


def format_recent_reference(row: pd.Series) -> str:
    author_items = [x.strip() for x in str(row["authors"]).split(";") if x.strip()]
    formatted = []
    for item in author_items[:10]:
        bits = item.split()
        if len(bits) > 1:
            formatted.append(f"{bits[0]}, {' '.join(bits[1:])}")
        else:
            formatted.append(item)
    authors = "; ".join(formatted)
    if len(author_items) > 10:
        authors += "; et al"
    title = clean_text(row["title"]).rstrip(".")
    journal = clean_text(row["journal_abbrev"])
    year = str(row["year"])
    volume = "" if pd.isna(row["volume"]) else clean_text(row["volume"])
    pages = "" if pd.isna(row["pages_or_elocation"]) else clean_text(row["pages_or_elocation"])
    doi = "" if pd.isna(row["doi"]) else clean_text(row["doi"])
    details = ", ".join(x for x in [year, volume, pages] if x)
    suffix = f" https://doi.org/{doi}." if doi else f" PMID: {row['pmid']}."
    return f"{authors}. {title}. {journal} {details}.{suffix}"


def load_references() -> list[str]:
    library = pd.read_csv(REFERENCE_LIBRARY, dtype={"pmid": str})
    by_pmid = library.set_index("pmid")
    refs = list(CLASSIC_REFERENCES[:8])
    ordered_groups = [
        ["39298548", "36142787", "26768243", "39113797", "32303283", "33073587", "32571916", "39110788", "31000567", "35922662", "38126209", "36229672", "33411069"],
        ["CLASSIC_ATLAS"],
        ["36932062", "33176333", "35491858", "36265491", "38513647", "39016438", "35709763", "38351505", "37130011"],
        ["CLASSIC_TIMP"],
        ["35513123"],
        ["CLASSIC_TGFB"],
        ["35064106", "38977708", "39572154", "38609039", "34560077", "38086793"],
    ]
    for group in ordered_groups:
        for item in group:
            if item == "CLASSIC_ATLAS":
                refs.append(CLASSIC_REFERENCES[9])
            elif item == "CLASSIC_TIMP":
                refs.append(CLASSIC_REFERENCES[8])
            elif item == "CLASSIC_TGFB":
                refs.append(CLASSIC_REFERENCES[10])
            elif item in by_pmid.index:
                refs.append(format_recent_reference(by_pmid.loc[item]))
    refs.extend(DATASET_REFERENCES)
    return refs


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def section_paragraphs(source: Document, start: str, end: str | None) -> list[tuple[str, str]]:
    collecting = False
    items: list[tuple[str, str]] = []
    for paragraph in source.paragraphs:
        if paragraph.text == start:
            collecting = True
        if collecting:
            if end and paragraph.text == end:
                break
            items.append((paragraph.style.name, paragraph.text))
    return items


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.25))
    legend = doc.add_paragraph(caption)
    legend.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in legend.runs:
        run.font.size = Pt(9)


def add_section_content(
    doc: Document,
    items: list[tuple[str, str]],
    citation_updates: dict[str, str],
    extra_after: dict[str, list[str]] | None = None,
    figures_after: dict[str, tuple[Path, str]] | None = None,
) -> None:
    extra_after = extra_after or {}
    figures_after = figures_after or {}
    for style, original in items:
        text_value = citation_updates.get(original, original)
        if not text_value.strip():
            continue
        if style.startswith("Heading"):
            level = int(style.split()[-1])
            if text_value == "Methods":
                text_value = "Materials and Methods"
            doc.add_heading(text_value, level=level)
        else:
            paragraph = doc.add_paragraph(text_value)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if original.startswith("Table 1. Public human kidney transcriptomic cohorts"):
                add_table(
                    doc,
                    ["Cohort", "Role", "Disease samples", "Control samples", "Tissue/platform", "Primary analytic use"],
                    [
                        ["GSE139061", "Discovery", "39 AKI", "9 reference nephrectomies", "Human kidney biopsy RNA-seq", "Expression, signatures, module discovery"],
                        ["GSE30718", "Discovery", "28 AKI/transplant injury", "11 controls", "Affymetrix microarray", "Expression, signatures, module discovery"],
                        ["GSE66494", "Discovery", "53 CKD/fibrosis", "8 controls", "Agilent microarray", "Expression, signatures, module discovery"],
                        ["GSE180394", "External validation", "44 kidney disease", "9 living donors; 6 tumor-nephrectomy sensitivity controls", "Microdissected tubules; Affymetrix", "External validation and robustness"],
                    ],
                )
        for extra_text in extra_after.get(original, []):
            paragraph = doc.add_paragraph(extra_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if original in figures_after:
            image_path, caption = figures_after[original]
            add_figure(doc, image_path, caption)


def build_manuscript() -> None:
    source = Document(SOURCE_DOCX)
    refs = load_references()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = None
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Yanzhao Ji 1,* and Zhihong Gao 1").bold = True
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.add_run(f"1 {AFFILIATION}")
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corr.add_run(f"* Correspondence: jiyanzhao@sxbqeh.com.cn; {ADDRESS}")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Maladaptive repair after kidney injury involves persistent tubular stress, inflammation, "
        "senescence, and extracellular matrix remodeling. We performed an integrative, platform-specific "
        "transcriptomic association study using three discovery cohorts and an independent microdissected "
        "tubular cohort. TIMP1 was directionally elevated in all discovery cohorts and was higher in "
        "GSE180394 disease samples than in living-donor controls (Hedges' g = 1.569, bootstrap 95% CI "
        "1.110-2.240). Twenty-five of 27 discovery disease-only TIMP1-signature correlations and all nine "
        "external correlations were false-discovery-rate significant. A 421-gene disease-only program score was "
        "externally reproduced with 401 detected genes (rho = 0.850); its correlation exceeded 1,000 "
        "expression- and variability-matched random programs (empirical P = 0.001) and remained positive "
        "after leaving out each discovery cohort. Leave-one-signature-out adjustment for an injury "
        "principal component "
        "attenuated many pathway associations, indicating substantial shared injury-state variance. "
        "Thus, TIMP1 is associated with a reproducible maladaptive injury-response context, but the "
        "cross-sectional data do not establish diagnostic utility, kidney specificity, prediction, or causality."
    )
    doc.add_paragraph(
        "Keywords: TIMP1; acute kidney injury; chronic kidney disease; tubular injury; "
        "maladaptive repair; extracellular matrix remodeling; fibrosis; transcriptomics"
    )

    citation_updates = {
        source.paragraphs[15].text: source.paragraphs[15].text,
        source.paragraphs[16].text: source.paragraphs[16].text.replace("[8]", "[8,12,15-19]"),
        source.paragraphs[17].text: source.paragraphs[17].text + " Recent single-cell and spatial atlases further show that injury-associated epithelial, immune, endothelial, and stromal states are spatially and molecularly heterogeneous [22-31].",
        source.paragraphs[18].text: source.paragraphs[18].text.replace("[9]", "[32,33]") + " Experimental evidence linking TIMP1 expression with kidney fibrosis susceptibility supports investigation of this association while not proving direct causality [33].",
        source.paragraphs[19].text: (
            "We therefore investigated the transcriptomic context of TIMP1 using three discovery "
            "bulk cohorts and one independent external cohort of microdissected human kidney tubules. "
            "We evaluated TIMP1 expression, nine prespecified injury-repair and ECM-related signatures, "
            "a discovery-derived disease-only program score, and a 12-gene core set. Robustness was "
            "examined by varying diagnosis and control definitions, matched-random-program testing, "
            "leave-one-discovery-cohort reconstruction, and leave-one-signature-out injury-PC adjustment. "
            "Single-cell findings were not used as inferential evidence because a complete multi-donor "
            "dataset was not locally reproducible. The overall workflow is summarized in Figure 1."
        ),
        source.paragraphs[22].text: (
            "This study was designed as an integrative transcriptomic association analysis rather than "
            "a systematic review. Candidate public datasets were identified from the project inventory "
            "and a targeted GEO search completed on 14 June 2026 using the exact query: "
            '(human[Organism]) AND (kidney OR renal) AND (AKI OR "acute kidney injury" OR CKD OR fibrosis) '
            'AND ("expression profiling by array"[DataSet Type] OR "expression profiling by high throughput '
            'sequencing"[DataSet Type]). Eligibility required '
            "a processed expression matrix, sample-level disease annotations, an identifiable reference "
            "group, and sufficient TIMP1/signature coverage. Cohorts were analyzed independently without "
            "pooling absolute expression across platforms. Candidate accessions and reasons for inclusion "
            "or exclusion are reported for the focused project candidates in Supplementary Table S1; the "
            "complete 268-record search-hit inventory is supplied separately for transparency [41-45]."
        ),
        source.paragraphs[23].text: (
            "The discovery analysis included GSE139061, GSE30718, and GSE66494 [41-43]. GSE180394 was "
            "analyzed independently as an external microdissected tubular validation cohort, with 44 "
            "kidney-disease samples and nine healthy living-donor controls as the primary comparison "
            "[44]. Six unaffected tumor-nephrectomy samples were reserved for control sensitivity analyses. "
            "GSE210622 was excluded from all inferential and exploratory manuscript results because the "
            "locally reconstructed object represented only one donor, whereas the original study was "
            "multi-donor [45]. The local single-donor files were retained only as historical audit material "
            "and were not summarized, plotted, tested, or cited as study evidence. "
            "Cohort roles and analytic uses are summarized in Table 1."
        ),
        source.paragraphs[26].text: (
            "Discovery matrices were processed separately. For GSE139061, the processed count matrix was "
            "transformed as log2(x + 1), quantile-normalized, mapped to gene symbols, and duplicate symbols "
            "were averaged. GSE30718 and GSE66494 series matrices were log2-transformed when required, "
            "quantile-normalized, annotated with GPL570 and GPL6480, respectively, and duplicate probes "
            "mapping to the same symbol were averaged. Ambiguous multi-symbol annotations retained the first "
            "listed symbol. Matrices contained 20,139, 21,755, and 19,553 genes, respectively. Sample labels "
            "and exclusions are documented in Supplementary Table S1."
        ),
        source.paragraphs[27].text: source.paragraphs[27].text + " The complete preprocessing manifest, input identifiers, and feature counts are supplied in Supplementary Table S1.",
        source.paragraphs[29].text: source.paragraphs[29].text + " Hedges' g confidence intervals for every reported expression contrast were estimated using one prespecified procedure: 2,000 stratified nonparametric bootstrap resamples, with disease and control groups resampled independently, percentile 95% intervals, and a recorded deterministic seed for each contrast.",
        source.paragraphs[31].text: (
            "Nine biologically motivated gene sets were frozen before external validation: ECM remodeling, "
            "collagen formation, TGF-beta signaling, inflammation, tubular injury, maladaptive repair, cellular "
            "senescence, fibrosis, and immune activation. They were compact literature-informed analysis sets, "
            "not unmodified database exports. The biological rationale, provenance category, version date, full "
            "membership, and overlap audit are provided in Supplementary Table S3. Within each dataset, per-gene "
            "expression was standardized across samples and signature scores were the mean of available genes. "
            "At least two genes were required, and TIMP1 was excluded to prevent self-correlation."
        ),
        source.paragraphs[33].text: source.paragraphs[33].text + " Spearman rho confidence intervals were estimated with 2,000 bootstrap resamples. Signature intercorrelations were quantified to make biological redundancy explicit.",
        source.paragraphs[35].text: source.paragraphs[35].text + " TIMP1 itself was excluded from module membership.",
        source.paragraphs[39].text: source.paragraphs[39].text.replace("399 of 421", "401 of 421"),
        source.paragraphs[43].text: source.paragraphs[43].text + " To evaluate shared injury-state confounding without mathematical self-inclusion, each focal signature was omitted in turn before PC1 construction from the remaining eight signatures. TIMP1 and the focal signature were residualized against this leave-one-signature-out injury PC1 before partial Spearman testing. PC1 variance explained and loadings were retained.",
        source.paragraphs[44].text: "",
        source.paragraphs[45].text: "",
        source.paragraphs[47].text: source.paragraphs[47].text + " Bootstrap analyses used 2,000 resamples; the matched-random-program benchmark used 1,000 repeats matched by mean-expression and variability deciles. OpenAI Codex (accessed 13-14 June 2026) was used for language editing, code drafting, and consistency checks; authors executed and verified all analyses. Random seeds are recorded in the analysis script.",
        source.paragraphs[50].text: (
            "TIMP1 expression was higher in disease samples in all three discovery cohorts when each dataset "
            "was analyzed independently (Figure 2). Mean disease-control differences were 1.086, 1.062, and "
            "0.417 in GSE139061, GSE30718, and GSE66494. Hedges' g values were 0.415 (bootstrap 95% CI "
            "-0.397 to 1.410), 1.293 (0.704 to 2.080), and 0.449 (0.168 to 0.747), respectively; BH-adjusted "
            "P values were 0.378, 3.26 x 10^-4, and 9.05 x 10^-3. Direction was consistent, but GSE139061 "
            "was imprecise and not statistically significant."
        ),
        source.paragraphs[51].text: (
            "In GSE180394, TIMP1 was higher in 44 disease samples than in nine living-donor controls "
            "(mean difference = 1.318; Hedges' g = 1.569; unadjusted two-sided Wilcoxon P = "
            "8.01 x 10^-5; bootstrap 95% CI for Hedges' g, 1.110-2.240; Figure 2). This was the "
            "prespecified primary external expression contrast; no multiplicity-adjusted P value is "
            "reported for this single primary comparison."
        ),
        source.paragraphs[55].text: "A reproducible disease-only program score is externally validated",
        source.paragraphs[53].text: source.paragraphs[53].text + " Bootstrap confidence intervals and exact adjusted P values are reported in Supplementary Table S3.",
        source.paragraphs[54].text: source.paragraphs[54].text + " All bootstrap rho confidence intervals remained positive.",
        source.paragraphs[57].text: (
            "In GSE180394, 401 of 421 stringent genes were detected. The program score correlated with TIMP1 "
            "at rho = 0.879 across all samples and rho = 0.850 among disease samples (Figure 4). The observed "
            "disease-only correlation exceeded 1,000 random programs matched on gene-level mean expression and "
            "variability using the same 53-sample standardization scope (random mean rho = 0.191; "
            "95th percentile = 0.491; empirical P = 0.001; Figure 6)."
        ),
        source.paragraphs[63].text: source.paragraphs[63].text + " In a stricter analysis, the focal signature was excluded before constructing an injury PC1 from the remaining eight signatures. These PC1s explained 72.6%-77.5% of external signature variance, and none of the nine external adjusted associations reached FDR < 0.05. This indicates that much of the observed co-variation is shared with a broad injury-state axis rather than representing independent TIMP1-specific pathway relationships.",
        source.paragraphs[64].text: "",
        source.paragraphs[65].text: "",
        source.paragraphs[67].text: (
            "In this integrative transcriptomic analysis, TIMP1 was reproducibly associated with a "
            "coordinated kidney injury-response context across four human datasets. TIMP1 correlated "
            "with nine prespecified programs, and a 421-gene discovery-derived score was externally "
            "reproduced and exceeded matched random programs. Leave-one-discovery-cohort scores "
            "remained correlated externally, but low gene-set overlap in two iterations shows that "
            "the exact module composition is not conserved. The evidence therefore supports a "
            "reproducible injury-state score rather than a fixed TIMP1-specific gene module."
        ),
        source.paragraphs[68].text: source.paragraphs[68].text.replace("[1-4]", "[1-4,9-21]").replace("[6,7]", "[6,7,15-21]"),
        source.paragraphs[69].text: source.paragraphs[69].text.replace("[9]", "[32]"),
        source.paragraphs[71].text: source.paragraphs[71].text + " Comparable compartment-resolved kidney atlases demonstrate the value of validating such programs across epithelial and stromal contexts [22-31].",
        source.paragraphs[72].text: source.paragraphs[72].text + " The leave-one-signature-out injury-PC sensitivity analysis further showed that disease-only restriction does not remove shared pathway severity or cell-composition structure; attenuation after this adjustment is therefore an important limit on TIMP1-specific interpretation.",
        source.paragraphs[73].text: (
            "Several limitations merit explicit acknowledgment. First, all cohorts were cross-sectional; "
            "no dataset followed the same individuals from AKI to CKD, and longitudinal claims are not "
            "supported. Second, systematic sample-level eGFR, histologic fibrosis, treatment, and outcome "
            "data were unavailable for covariate adjustment. Third, bulk transcriptomic associations may "
            "reflect cell-composition shifts as well as within-cell-state changes. Fourth, GSE180394 had "
            "nine living-donor controls and an etiologically heterogeneous disease group. Fifth, the compact "
            "signatures were literature-informed and partly overlapping rather than independent ontology-derived "
            "pathways; their provenance and overlap are audited in Supplementary Table S3. Sixth, complete "
            "Hallmark, GO, Reactome, and KEGG enrichment could not be computed in the current environment, "
            "so available enrichment is exploratory. Seventh, leave-one-discovery-cohort analysis showed "
            "stable external score correlations but unstable gene membership; the 421 genes are not a "
            "compositionally conserved module. Eighth, GSE210622 was excluded from all inferential and "
            "exploratory manuscript results because only one donor was locally reconstructed. Finally, no "
            "protein-level validation, spatial co-localization, longitudinal validation, or experimental "
            "perturbation was performed."
        ),
        source.paragraphs[74].text: source.paragraphs[74].text.replace("[10]", "[22-31]"),
        source.paragraphs[76].text: (
            "TIMP1 is directionally elevated across human kidney injury and fibrosis datasets and is associated "
            "with tubular injury-repair, ECM remodeling, inflammatory activation, senescence, fibrosis-related "
            "programs, and a reproducible disease-only transcriptomic score. External tubular validation and "
            "a matched-random-program benchmark strengthen the association-based evidence, while low "
            "leave-one-cohort gene-set overlap argues against a fixed conserved module. Attenuation after "
            "leave-one-signature-out injury-PC adjustment supports a conservative "
            "interpretation: TIMP1 is a candidate component or readout of a wider maladaptive repair and wound-"
            "response state. Longitudinal, spatial, protein-level, and perturbation studies are required before "
            "claims regarding prediction, diagnostic performance, kidney specificity, or causality."
        ),
    }

    extra_after = {
        source.paragraphs[23].text: [
            "The complete GEO hit inventory is provided for search transparency, but the search was not "
            "conducted or reported as a systematic review. Detailed eligibility decisions are provided for "
            "the focused candidate accessions evaluated in the project workflow. Other search hits are "
            "identified as not included after title/summary and eligibility screening without implying that "
            "a PRISMA-style full-text exclusion reason was assigned to every record."
        ],
        source.paragraphs[35].text: [
            "Program-score specificity was evaluated in two additional ways. First, the gene set was re-derived after leaving out each discovery cohort and then scored in GSE180394. Second, 1,000 random gene programs were matched to the observed external program by mean-expression and variability deciles. All external program scores were standardized across the same 53-sample cohort before disease-only correlation. These analyses tested dependence on a single discovery cohort and whether a similarly sized, similarly expressed gene set would show comparable correlation by chance."
        ],
        source.paragraphs[57].text: [
            "Leave-one-discovery-cohort-out scores remained positively associated with TIMP1 in GSE180394: rho = 0.854 after omitting GSE139061, 0.857 after omitting GSE30718, and 0.789 after omitting GSE66494. Program sizes differed substantially (13-403 genes), and Jaccard overlap with the full 421-gene set was only 0.031 and 0.069 in two iterations. Thus, score-level association was reproducible, whereas exact gene membership was not."
        ],
        source.paragraphs[72].text: [
            "A shared-injury or cell-composition explanation remains plausible. TIMP1 and the nine signatures can rise together because samples contain different proportions of injured epithelium, inflammatory cells, activated stroma, or globally more severe tissue damage. Marked attenuation after leave-one-signature-out injury-PC adjustment supports this explanation. Because this PC is a transcriptomic proxy rather than a measured clinical severity variable, the analysis is a stringent sensitivity test and not proof that true injury severity or cell composition has been fully controlled."
        ],
    }
    figures_after = {
        source.paragraphs[19].text: (
            ROOT / "figures/timp1_validation/v6_manuscript_figure_1_workflow_cohort_overview.png",
            "Figure 1. Study workflow and cohort roles. Discovery cohorts were analyzed independently, followed by external tubular validation and prespecified robustness analyses."
        ),
        source.paragraphs[51].text: (
            ROOT / "figures/timp1_validation/v5_manuscript_figure_2_TIMP1_expression_cohorts.png",
            "Figure 2. TIMP1 expression across discovery and external cohorts. Points show samples; boxes summarize group distributions. Exact statistics and bootstrap confidence intervals are in Supplementary Table S2."
        ),
        source.paragraphs[54].text: (
            ROOT / "figures/timp1_validation/v5_manuscript_figure_3_signature_correlation_heatmap.png",
            "Figure 3. Spearman correlations between TIMP1 and nine prespecified signatures. TIMP1 was excluded from signature scores. Stars denote BH-adjusted significance."
        ),
        source.paragraphs[57].text: (
            ROOT / "figures/timp1_validation/v6_manuscript_figure_4_GSE180394_program_score_scatter.png",
            "Figure 4. External validation of the stringent discovery-derived program score in GSE180394. The score used 401 detected genes."
        ),
        source.paragraphs[59].text: (
            ROOT / "figures/timp1_validation/v5_manuscript_figure_5_core_gene_correlation_heatmap.png",
            "Figure 5. External disease-only correlations between TIMP1 and the 12 cross-cohort core genes. Stars denote BH-adjusted significance."
        ),
        source.paragraphs[57].text + "__random": (
            ROOT / "figures/timp1_validation/v6_module_random_matched_benchmark.png",
            "Figure 6. Matched-random-program benchmark. The observed GSE180394 disease-only score correlation is compared with 1,000 programs matched by mean expression and variability. All scores were standardized across the same 53-sample cohort before disease-only correlation."
        ),
    }
    random_figure = figures_after.pop(source.paragraphs[57].text + "__random")

    add_section_content(doc, section_paragraphs(source, "Introduction", "Methods"), citation_updates, extra_after, figures_after)
    add_section_content(doc, section_paragraphs(source, "Results", "Discussion"), citation_updates, extra_after, figures_after)
    add_figure(doc, *random_figure)
    add_section_content(doc, section_paragraphs(source, "Discussion", "Conclusions"), citation_updates, extra_after, figures_after)
    add_section_content(doc, section_paragraphs(source, "Methods", "Results"), citation_updates, extra_after, figures_after)
    add_section_content(doc, section_paragraphs(source, "Conclusions", "Declarations"), citation_updates, extra_after, figures_after)

    doc.add_heading("Supplementary Materials", level=1)
    doc.add_paragraph(
        "The following supporting information can be downloaded with the article: "
        "Table S1, cohort characteristics and analytic roles; Table S2, TIMP1 expression "
        "statistics; Table S3, prespecified signatures and TIMP1-signature correlations; "
        "Table S4, program-score and core-gene validation; Table S5, robustness analyses, feature "
        "coverage, and missing-data log."
    )
    doc.add_heading("Author Contributions", level=1)
    doc.add_paragraph(
        "Conceptualization, Y.J.; methodology, Y.J. and Z.G.; software, Y.J.; validation, "
        "Z.G.; formal analysis, Y.J.; investigation, Y.J. and Z.G.; data curation, Y.J. and "
        "Z.G.; visualization, Y.J.; writing-original draft preparation, Y.J.; writing-review "
        "and editing, Y.J. and Z.G.; supervision, Y.J.; project administration, Y.J.; funding "
        "acquisition, Y.J. All authors have read and agreed to the published version of the manuscript."
    )
    doc.add_heading("Funding", level=1)
    doc.add_paragraph(
        "This research was supported by the Fundamental Research Program of Shanxi Province "
        "(No. 202403021222407) and the Shanxi Bethune Hospital Project (No. 2023RC06). "
        "The funders had no role in study design, data collection, analysis, interpretation, "
        "manuscript preparation, or the decision to submit the work."
    )
    doc.add_heading("Institutional Review Board Statement", level=1)
    doc.add_paragraph(
        "Ethical review and approval were waived for this study because it was a secondary "
        "analysis of publicly available, de-identified transcriptomic datasets and involved "
        "no new participant recruitment or biospecimen collection."
    )
    doc.add_heading("Informed Consent Statement", level=1)
    doc.add_paragraph("Not applicable.")
    doc.add_heading("Data Availability Statement", level=1)
    doc.add_paragraph(
        "The public transcriptomic datasets are available from the NCBI Gene Expression "
        "Omnibus under accession numbers GSE139061, GSE30718, GSE66494, and GSE180394. "
        "The current public repository and Zenodo concept record contain the previously released "
        f"analysis snapshot ({GITHUB_URL}; {ZENODO_DOI_URL}). The revised v1.2 analysis package, "
        "including the unified bootstrap results and current source files, is prepared locally and "
        "will be released and linked in the final manuscript before journal submission."
    )
    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph(
        "The authors thank the investigators and participants of the public GEO studies "
        "analyzed in this work."
    )
    doc.add_heading("Conflicts of Interest", level=1)
    doc.add_paragraph("The authors declare no conflict of interest.")
    doc.add_heading("Use of Artificial Intelligence", level=1)
    doc.add_paragraph(
        "OpenAI Codex (accessed 13-14 June 2026) was used for English-language editing, "
        "code drafting, and consistency checks. It was not used to make autonomous scientific "
        "or clinical decisions. The authors inspected and executed the analysis code, verified "
        "numerical results against source outputs, checked references against bibliographic records, "
        "and take full responsibility for the manuscript."
    )

    doc.add_heading("References", level=1)
    for index, reference in enumerate(refs, start=1):
        paragraph = doc.add_paragraph(f"{index}. {reference}")
        paragraph.paragraph_format.first_line_indent = Inches(-0.2)
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(3)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("TIMP1 integrative transcriptomic analysis - IJMS submission draft")
    doc.save(OUT / "TIMP1_IJMS_Manuscript.docx")


def dataframe_with_source(df: pd.DataFrame, source: str) -> pd.DataFrame:
    output = df.copy()
    output.insert(0, "source_file", source)
    return output


def format_statistical_zeros(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()
    for column in output.columns:
        name = str(column).lower()
        is_count = any(
            token in name
            for token in ["count", "iterations", "repetitions", "genes"]
        )
        is_p_value = (
            any(token in name for token in ["p_value", "pvalue", "p_adj", "padj"])
            or name in {"p", "fdr"}
        ) and not is_count
        if is_p_value:
            output[column] = output[column].map(
                lambda value: "<2.2e-16"
                if isinstance(value, (int, float)) and not pd.isna(value) and value == 0
                else value
            )
    return output


def build_supplement() -> dict[str, list[tuple[str, pd.DataFrame]]]:
    cohorts = pd.DataFrame(
        [
            ["GSE139061", "Discovery", "39 AKI", "9 controls", "Kidney biopsy RNA-seq", "Expression, signatures, module discovery"],
            ["GSE30718", "Discovery", "28 AKI/transplant injury", "11 controls", "Affymetrix microarray", "Expression, signatures, module discovery"],
            ["GSE66494", "Discovery", "53 CKD/fibrosis", "8 controls", "Agilent microarray", "Expression, signatures, module discovery"],
            ["GSE180394", "External validation", "44 kidney disease", "9 living donors; 6 tumor-nephrectomy sensitivity controls", "Microdissected tubules; Affymetrix", "External validation and robustness"],
        ],
        columns=["dataset", "role", "disease_samples", "control_samples", "tissue_platform", "analytic_use"],
    )
    screening = pd.DataFrame(
        [
            ["GSE139061", "Included", "Discovery bulk cohort", "Processed matrix, disease labels, reference nephrectomy samples, and required gene coverage available"],
            ["GSE30718", "Included", "Discovery bulk cohort", "Processed matrix, transplant injury labels, control samples, and platform annotation available"],
            ["GSE66494", "Included", "Discovery bulk cohort", "Processed matrix, CKD/fibrosis labels, controls, and platform annotation available"],
            ["GSE180394", "Included", "External tubular validation", "Independent microdissected tubule cohort with living-donor and sensitivity controls"],
            ["GSE210622", "Excluded from inference", "Screened single-cell candidate", "Only one donor was locally reproducible; the original multi-donor study could not be reconstructed from the available local files without selective donor use"],
            ["GSE133288", "Excluded", "Screened candidate", "No locally reproducible processed matrix and matched annotation meeting the predefined workflow requirements"],
            ["GSE267242", "Excluded", "Screened candidate", "Available files and study context did not meet the predefined human kidney cohort requirements for this analysis"],
        ],
        columns=["accession", "decision", "role", "reason"],
    )
    preprocessing = pd.DataFrame(
        [
            ["GSE139061", "Processed count matrix", "log2(x + 1); quantile normalization", "Gene symbols; duplicate symbols averaged", 20139],
            ["GSE30718", "GEO series matrix", "Conditional log2; quantile normalization", "GPL570; first symbol for ambiguous probes; duplicates averaged", 21755],
            ["GSE66494", "GEO series matrix", "Conditional log2; quantile normalization", "GPL6480; first symbol for ambiguous probes; duplicates averaged", 19553],
            ["GSE180394", "Processed GEO expression object", "As supplied; analyzed separately", "GPL19983/NCBI gene_info; duplicate symbols averaged", 24845],
        ],
        columns=["dataset", "input", "transformation", "annotation_and_collapse", "final_gene_count"],
    )
    search_audit = pd.DataFrame(
        [
            ["Database", "NCBI Gene Expression Omnibus via Entrez GDS"],
            ["Search date", "14 June 2026"],
            [
                "Query",
                '(human[Organism]) AND (kidney OR renal) AND (AKI OR "acute kidney injury" OR CKD OR fibrosis) AND ("expression profiling by array"[DataSet Type] OR "expression profiling by high throughput sequencing"[DataSet Type])',
            ],
            ["Records returned", "268"],
            [
                "Screening approach",
                "Complete hit inventory plus detailed eligibility audit for focused project candidates; not a systematic review and not a PRISMA full-text exclusion log",
            ],
        ],
        columns=["field", "value"],
    )
    s1 = [
        ("Cohort characteristics", cohorts),
        ("Candidate dataset screening audit", screening),
        ("Reproducible GEO search audit", search_audit),
        (
            "Complete GEO search-hit inventory",
            pd.read_csv(
                ROOT
                / "results/timp1_validation/geo_search_hits_2026-06-14.csv"
            ),
        ),
        ("Preprocessing manifest", preprocessing),
    ]

    bulk = pd.read_csv(ROOT / "results/timp1_validation/TIMP1_bulk_validation_statistics.csv")
    external = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_TIMP1_group_statistics.csv")
    expression_unified = pd.read_csv(
        ROOT
        / "results/timp1_validation/TIMP1_expression_statistics_unified_bootstrap_v7.csv"
    )
    s2 = [
        (
            "Authoritative expression statistics and unified bootstrap confidence intervals",
            dataframe_with_source(
                expression_unified,
                "TIMP1_expression_statistics_unified_bootstrap_v7.csv",
            ),
        ),
    ]

    gene_sets = pd.read_csv(ROOT / "results/timp1_validation/signature_gene_sets_used.csv")
    correlations = pd.read_csv(ROOT / "results/timp1_validation/TIMP1_signature_correlations.csv")
    external_corr = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_signature_correlations.csv")
    signature_provenance = pd.read_csv(ROOT / "results/timp1_validation/signature_provenance_v5.csv")
    overlap = pd.read_csv(ROOT / "results/timp1_validation/signature_overlap_audit_v5.csv")
    intercorrelations = pd.read_csv(ROOT / "results/timp1_validation/signature_intercorrelations_disease_only_v5.csv")
    correlation_ci = pd.read_csv(ROOT / "results/timp1_validation/TIMP1_signature_correlation_ci_v5.csv")
    broad_adjustment = pd.read_csv(
        ROOT
        / "results/timp1_validation/signature_leave_one_out_injury_pc_adjustment_v6.csv"
    )
    injury_pc_loadings = pd.read_csv(
        ROOT
        / "results/timp1_validation/signature_leave_one_out_injury_pc_loadings_v6.csv"
    )
    s3 = [
        ("Signature provenance", dataframe_with_source(signature_provenance, "signature_provenance_v5.csv")),
        ("Prespecified signature membership", dataframe_with_source(gene_sets, "signature_gene_sets_used.csv")),
        ("Signature overlap audit", dataframe_with_source(overlap, "signature_overlap_audit_v5.csv")),
        ("Disease-only signature intercorrelations", dataframe_with_source(intercorrelations, "signature_intercorrelations_disease_only_v5.csv")),
        ("Discovery correlations", dataframe_with_source(correlations, "TIMP1_signature_correlations.csv")),
        ("External correlations", dataframe_with_source(external_corr, "external_GSE180394_signature_correlations.csv")),
        ("Bootstrap correlation confidence intervals", dataframe_with_source(correlation_ci, "TIMP1_signature_correlation_ci_v5.csv")),
        ("Leave-one-signature-out injury-PC adjustment", dataframe_with_source(broad_adjustment, "signature_leave_one_out_injury_pc_adjustment_v6.csv")),
        ("Leave-one-signature-out injury-PC loadings", dataframe_with_source(injury_pc_loadings, "signature_leave_one_out_injury_pc_loadings_v6.csv")),
    ]

    core = pd.read_csv(ROOT / "results/timp1_validation/core_module_genes_summary.csv")
    stringent = pd.read_csv(ROOT / "results/timp1_validation/stringent_TIMP1_correlated_module.csv")
    module_external = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_stringent_module_correlations.csv")
    core_external = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_core_module_correlations.csv")
    loco = pd.read_csv(
        ROOT / "results/timp1_validation/module_leave_one_discovery_out_v6.csv"
    )
    random_benchmark = pd.read_csv(
        ROOT / "results/timp1_validation/module_random_matched_benchmark_v6.csv"
    )
    s4 = [
        ("Core module summary", dataframe_with_source(core, "core_module_genes_summary.csv")),
        ("Stringent 421-gene module", dataframe_with_source(stringent, "stringent_TIMP1_correlated_module.csv")),
        ("External module validation", dataframe_with_source(module_external, "external_GSE180394_stringent_module_correlations.csv")),
        ("External core-gene validation", dataframe_with_source(core_external, "external_GSE180394_core_module_correlations.csv")),
        ("Leave-one-discovery-cohort-out validation", dataframe_with_source(loco, "module_leave_one_discovery_out_v6.csv")),
        ("Matched-random-program benchmark", dataframe_with_source(random_benchmark, "module_random_matched_benchmark_v6.csv")),
    ]

    lodo = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_sensitivity_leave_one_diagnosis_out_summary.csv")
    robust = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_sensitivity_robust_correlations.csv")
    coverage = pd.read_csv(ROOT / "results/timp1_validation/external_GSE180394_input_feature_coverage.csv")
    missing = pd.read_csv(ROOT / "results/timp1_validation/missing_data_log.csv")
    missing["record_status"] = "Current"
    historical_singlecell = missing["analysis_step"].eq("singlecell_validation")
    missing.loc[historical_singlecell, "record_status"] = (
        "Historical audit; superseded by singlecell_exclusion_policy_v7.csv"
    )
    missing.loc[historical_singlecell, "action_taken"] = (
        "Historical action superseded. Current manuscript policy excludes these "
        "single-cell inputs from all inferential and exploratory results."
    )
    singlecell_policy = pd.read_csv(
        ROOT / "results/timp1_validation/singlecell_exclusion_policy_v7.csv"
    )
    s5 = [
        ("Leave-one-diagnosis-out summary", dataframe_with_source(lodo, "external_GSE180394_sensitivity_leave_one_diagnosis_out_summary.csv")),
        ("Diagnosis-adjusted robustness", dataframe_with_source(robust, "external_GSE180394_sensitivity_robust_correlations.csv")),
        ("External feature coverage", dataframe_with_source(coverage, "external_GSE180394_input_feature_coverage.csv")),
        (
            "Current single-cell exclusion policy",
            dataframe_with_source(
                singlecell_policy, "singlecell_exclusion_policy_v7.csv"
            ),
        ),
        ("Missing-data log", dataframe_with_source(missing, "missing_data_log.csv")),
    ]
    sections = {
        "S1_Cohorts": s1,
        "S2_Expression": s2,
        "S3_Signatures": s3,
        "S4_Modules": s4,
        "S5_Sensitivity": s5,
    }

    workbook = Workbook()
    workbook.remove(workbook.active)
    for sheet_name, tables in sections.items():
        sheet = workbook.create_sheet(sheet_name)
        sheet.sheet_view.showGridLines = False
        row = 1
        for title, frame in tables:
            frame = format_statistical_zeros(frame)
            sheet.cell(row=row, column=1, value=title)
            sheet.cell(row=row, column=1).font = Font(bold=True, size=13, color="FFFFFF")
            sheet.cell(row=row, column=1).fill = PatternFill("solid", fgColor="1F4E78")
            sheet.merge_cells(start_row=row, start_column=1, end_row=row, end_column=max(1, len(frame.columns)))
            row += 1
            for col_idx, column in enumerate(frame.columns, start=1):
                cell = sheet.cell(row=row, column=col_idx, value=str(column))
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="5B9BD5")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            row += 1
            for values in frame.itertuples(index=False, name=None):
                for col_idx, value in enumerate(values, start=1):
                    if pd.isna(value):
                        value = None
                    cell = sheet.cell(row=row, column=col_idx, value=value)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                row += 1
            row += 2
        sheet.freeze_panes = "A3"
        sheet.auto_filter.ref = sheet.dimensions
        for col_idx in range(1, sheet.max_column + 1):
            max_len = max(len(str(sheet.cell(r, col_idx).value or "")) for r in range(1, min(sheet.max_row, 250) + 1))
            sheet.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 11), 34)
    xlsx = OUT / "Supplementary_Tables_S1-S5.xlsx"
    workbook.save(xlsx)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("SuppTitle", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=14)
    cell_style = ParagraphStyle("Cell", parent=styles["BodyText"], fontSize=5.5, leading=6.3)
    pdf = SimpleDocTemplate(
        str(OUT / "Supplementary_Tables_S1-S5.pdf"),
        pagesize=landscape(A4),
        leftMargin=8 * mm,
        rightMargin=8 * mm,
        topMargin=8 * mm,
        bottomMargin=8 * mm,
    )
    story = []
    for sheet_name, tables in sections.items():
        story.append(Paragraph(sheet_name.replace("_", " "), title_style))
        for table_title, frame in tables:
            frame = format_statistical_zeros(frame)
            story.append(Paragraph(table_title, styles["Heading2"]))
            display = frame.copy()
            max_rows = 80
            if len(display) > max_rows:
                display = display.head(max_rows)
                story.append(Paragraph(f"PDF preview shows first {max_rows} rows; complete table is in the Excel file.", styles["Italic"]))
            data = [[Paragraph(clean_text(c), cell_style) for c in display.columns]]
            for values in display.astype(object).where(pd.notnull(display), "").itertuples(index=False, name=None):
                data.append([Paragraph(clean_text(v), cell_style) for v in values])
            usable = landscape(A4)[0] - 16 * mm
            widths = [usable / max(1, len(display.columns))] * len(display.columns)
            table = Table(data, colWidths=widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C9D6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F7FA")]),
                    ]
                )
            )
            story.append(table)
            story.append(Paragraph(" ", styles["BodyText"]))
        story.append(PageBreak())
    pdf.build(story)
    return sections


def build_cover_letter() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_paragraph("14 June 2026")
    doc.add_paragraph("Editorial Office\nInternational Journal of Molecular Sciences")
    doc.add_paragraph("Dear Editors,")
    doc.add_paragraph(
        f"We are pleased to submit the original article entitled \"{TITLE}\" for consideration "
        "in the International Journal of Molecular Sciences."
    )
    doc.add_paragraph(
        "This integrative transcriptomic study evaluates TIMP1 across three discovery kidney "
        "cohorts and an independent microdissected tubular cohort. The work emphasizes "
        "cross-dataset reproducibility, disease-only pathway associations, a reproducible disease-state "
        "program score, leave-one-discovery-cohort validation, a matched-random-program benchmark, and "
        "robustness to diagnosis, control definitions, and leave-one-signature-out injury-state factors. The manuscript treats "
        "TIMP1 as a candidate injury and extracellular-matrix-remodeling-associated gene and "
        "does not claim kidney specificity, diagnostic utility, or causality. Attenuation after "
        "leave-one-signature-out injury-PC adjustment and unstable leave-one-cohort gene membership "
        "are reported explicitly."
    )
    doc.add_paragraph(
        "The manuscript is original, is not under consideration elsewhere, and has been approved "
        "by both authors. The authors declare no conflict of interest. The analysis used publicly "
        "available de-identified datasets and did not involve new participant recruitment."
    )
    doc.add_paragraph("Thank you for considering our manuscript.")
    doc.add_paragraph(
        "Sincerely,\nYanzhao Ji, Corresponding Author\n"
        f"{AFFILIATION}\n{ADDRESS}\nEmail: jiyanzhao@sxbqeh.com.cn"
    )
    doc.save(OUT / "TIMP1_IJMS_Cover_Letter.docx")


def copy_tree_contents(source: Path, destination: Path, patterns: tuple[str, ...] | None = None) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    for item in source.iterdir():
        if item.is_file() and (patterns is None or item.suffix.lower() in patterns):
            shutil.copy2(item, destination / item.name)


def build_repository() -> None:
    # Preserve an initialized Git repository and refresh tracked content in place.
    REPO.mkdir(parents=True, exist_ok=True)
    ZENODO.mkdir(parents=True, exist_ok=True)

    copy_tree_contents(ROOT / "scripts", REPO / "scripts", (".r", ".py", ".ps1"))
    copy_tree_contents(ROOT / "results/timp1_validation", REPO / "results/timp1_validation")
    copy_tree_contents(ROOT / "figures/timp1_validation", REPO / "figures/timp1_validation", (".pdf", ".png"))
    copy_tree_contents(ROOT / "TIMP1_AKI_CKD_project/data/processed", REPO / "data/processed", (".gz", ".csv"))
    copy_tree_contents(ROOT / "TIMP1_AKI_CKD_project/data/metadata", REPO / "data/metadata", (".csv",))
    shutil.copy2(ROOT / "AGENTS.md", REPO / "ANALYSIS_GUARDRAILS.md")

    # GSE210622 was excluded from inferential analyses in the revised manuscript.
    for data_root in [REPO / "data", ZENODO / "data"]:
        if data_root.exists():
            for stale_path in data_root.rglob("GSE210622*"):
                if stale_path.is_file():
                    stale_path.unlink()

    # Keep historical single-donor audits in the working project, but exclude them
    # from the current publication evidence package to prevent status ambiguity.
    publication_exclusions = [
        "scripts/*singlecell*",
        "results/timp1_validation/singlecell*",
        "results/timp1_validation/manuscript_first_draft*",
        "results/timp1_validation/manuscript_abstract_options*",
        "results/timp1_validation/updated_validation_report*",
        "results/timp1_validation/sessionInfo_singlecell_audit.txt",
    ]
    for package_root in [REPO, ZENODO]:
        for pattern in publication_exclusions:
            for stale_path in package_root.glob(pattern):
                if stale_path.is_file():
                    stale_path.unlink()

        public_log = (
            package_root
            / "results/timp1_validation/missing_data_log.csv"
        )
        if public_log.exists():
            public_missing = pd.read_csv(public_log)
            historical = public_missing["analysis_step"].eq(
                "singlecell_validation"
            )
            public_missing.loc[historical, "action_taken"] = (
                "Historical audit superseded. Current manuscript policy excludes "
                "these inputs from all inferential and exploratory results."
            )
            public_missing["record_status"] = "Current"
            public_missing.loc[historical, "record_status"] = (
                "Historical audit; superseded by "
                "singlecell_exclusion_policy_v7.csv"
            )
            public_missing.to_csv(public_log, index=False, encoding="utf-8-sig")

    for folder in ["scripts", "results", "figures", "data"]:
        shutil.copytree(REPO / folder, ZENODO / folder, dirs_exist_ok=True)
    readme = f"""# TIMP1 kidney injury transcriptomic analysis

This repository contains the reproducible code, derived tables, intermediate
outputs, session information, source data, and figures for:

**{TITLE}**

Authors: Yanzhao Ji and Zhihong Gao.

The analysis positions TIMP1 as a candidate kidney injury and extracellular
matrix remodeling-associated gene. Results are association-based and do not
establish diagnostic utility, kidney specificity, longitudinal prediction, or
causality.

## Contents

- `scripts/`: R and Python workflows.
- `results/timp1_validation/`: statistical tables, intermediate results,
  session information, interpretations, and missing-data records.
- `figures/timp1_validation/`: PDF and PNG figure files.
- `data/processed/`: processed expression matrices small enough for GitHub.
- `data/metadata/`: sample metadata.

The complete versioned archive is deposited on Zenodo.

## Archived release

- GitHub release: {GITHUB_RELEASE_URL}
- Zenodo DOI: {ZENODO_DOI_URL}

## Public datasets

NCBI GEO: GSE139061, GSE30718, GSE66494, and GSE180394.

## Reproduction

R 4.6.0 was used for the principal analyses. Run the scripts in numeric or
workflow order from the repository root. Package versions are recorded in the
`sessionInfo*.txt` files under `results/timp1_validation/`.

## Missing data

Unavailable optional datasets or resources are recorded in
`results/timp1_validation/missing_data_log.csv`; affected units are skipped
without silently terminating the remaining workflow.
"""
    (REPO / "README.md").write_text(readme, encoding="utf-8")
    (ZENODO / "README.md").write_text(readme, encoding="utf-8")
    license_text = (
        "MIT License\n\nCopyright (c) 2026 Yanzhao Ji and Zhihong Gao\n\n"
        "Permission is hereby granted, free of charge, to any person obtaining a copy "
        "of this software and associated documentation files (the \"Software\"), to deal "
        "in the Software without restriction, including without limitation the rights "
        "to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies "
        "of the Software, subject to inclusion of this notice.\n\nTHE SOFTWARE IS PROVIDED "
        "\"AS IS\", WITHOUT WARRANTY OF ANY KIND."
    )
    (REPO / "LICENSE").write_text(license_text, encoding="utf-8")
    (ZENODO / "LICENSE").write_text(license_text, encoding="utf-8")
    citation = f"""cff-version: 1.2.0
message: "If you use this repository, please cite the archived release."
title: "{TITLE}"
type: software
authors:
  - family-names: Ji
    given-names: Yanzhao
    email: jiyanzhao@sxbqeh.com.cn
  - family-names: Gao
    given-names: Zhihong
version: 1.2.0
date-released: 2026-06-14
license: MIT
repository-code: "{GITHUB_URL}"
doi: "{ZENODO_DOI}"
"""
    (REPO / "CITATION.cff").write_text(citation, encoding="utf-8")
    (ZENODO / "CITATION.cff").write_text(citation, encoding="utf-8")
    zenodo_json = {
        "title": TITLE,
        "upload_type": "software",
        "description": "Reproducible code, derived data, intermediate outputs, and figures for an integrative TIMP1 kidney transcriptomic analysis.",
        "creators": [
            {"name": "Ji, Yanzhao", "affiliation": AFFILIATION},
            {"name": "Gao, Zhihong", "affiliation": AFFILIATION},
        ],
        "keywords": ["TIMP1", "acute kidney injury", "chronic kidney disease", "fibrosis", "transcriptomics"],
        "license": "mit",
        "version": "1.2.0",
        "publication_date": "2026-06-14",
    }
    (ZENODO / ".zenodo.json").write_text(json.dumps(zenodo_json, indent=2), encoding="utf-8")

    manifest_rows = []
    for path in sorted(ZENODO.rglob("*")):
        if path.is_file():
            manifest_rows.append(
                {
                    "relative_path": path.relative_to(ZENODO).as_posix(),
                    "size_bytes": path.stat().st_size,
                }
            )
    with (ZENODO / "FILE_MANIFEST.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=["relative_path", "size_bytes"])
        writer.writeheader()
        writer.writerows(manifest_rows)


def build_readme() -> None:
    text = f"""IJMS submission package
Prepared: 2026-06-14

Files
- TIMP1_IJMS_Manuscript.docx
- TIMP1_IJMS_Cover_Letter.docx
- Supplementary_Tables_S1-S5.xlsx
- Supplementary_Tables_S1-S5.pdf
- references_verified.csv
- Figure_1 to Figure_6 in PDF and PNG formats

Authors
- Yanzhao Ji: first and corresponding author
- Zhihong Gao
- Affiliation: {AFFILIATION}
- Correspondence: jiyanzhao@sxbqeh.com.cn

Repository status
- Current public GitHub repository: {GITHUB_URL}
- Current public GitHub release v1.1.0: {GITHUB_RELEASE_URL}
- Current Zenodo concept DOI: {ZENODO_DOI_URL}
- Revised v1.2.0 files are prepared locally but have not yet been pushed or archived.
- The complete pending local Zenodo payload is under
  public_repository/TIMP1_AKI_CKD_transcriptomics_zenodo_payload.
"""
    (OUT / "README_Submission_Files.txt").write_text(text, encoding="utf-8")


def build_revision_summary() -> None:
    text = f"""# Revision summary

Prepared: 14 June 2026

## Major revisions

1. Added five verified primary dataset citations and linked each analyzed GEO accession to its original publication or GEO record.
2. Added a candidate-dataset screening audit and dataset-specific preprocessing manifest to Supplementary Table S1.
3. Documented the literature-informed provenance, version date, membership, overlap, and disease-only intercorrelations of all nine signatures.
4. Added 2,000-resample bootstrap confidence intervals for TIMP1 expression effect sizes and TIMP1-signature correlations.
5. Reframed the 421-gene result as a reproducible disease-state program score rather than a compositionally conserved module; added leave-one-discovery-cohort-out validation and a 1,000-repeat expression/variability-matched random-program benchmark.
6. Added leave-one-signature-out injury-PC adjustment, including variance explained and loadings. The attenuation of associations is reported as evidence that shared injury severity and cell composition remain plausible alternative explanations.
7. Corrected external stringent-module coverage to 401 of 421 genes based on the actual scoring matrix.
8. Replaced violin plots with sample-level box plots, added significance marks to heatmaps, embedded six main figures in the manuscript, and supplied all figures as PDF and PNG.
9. Added a complete 268-record GEO search inventory, funder-role and artificial-intelligence-use statements, and removed non-analyzed accessions from the Data Availability Statement.
10. Removed single-donor single-cell results from inferential evidence because the original multi-donor dataset was not completely reconstructed locally.

## Reproducibility

New validation outputs are under `results/timp1_validation/`; new figures are under `figures/timp1_validation/`. A local repository payload has been prepared for version 1.2.0, but it has not yet been pushed or archived.

- GitHub: {GITHUB_URL}
- Zenodo DOI: {ZENODO_DOI_URL}
"""
    (OUT / "Revision_Summary.md").write_text(text, encoding="utf-8")


def copy_figures() -> None:
    mapping = {
        "v6_manuscript_figure_1_workflow_cohort_overview": "Figure_1_TIMP1_Validation_Workflow",
        "v5_manuscript_figure_2_TIMP1_expression_cohorts": "Figure_2_TIMP1_Expression_Cohorts",
        "v5_manuscript_figure_3_signature_correlation_heatmap": "Figure_3_TIMP1_Signature_Correlations",
        "v6_manuscript_figure_4_GSE180394_program_score_scatter": "Figure_4_TIMP1_Program_Score_GSE180394",
        "v5_manuscript_figure_5_core_gene_correlation_heatmap": "Figure_5_TIMP1_Core_Gene_Correlations",
        "v6_module_random_matched_benchmark": "Figure_6_TIMP1_Program_Random_Benchmark",
    }
    for source_name, target_name in mapping.items():
        for suffix in [".pdf", ".png"]:
            shutil.copy2(
                ROOT / "figures/timp1_validation" / f"{source_name}{suffix}",
                OUT / f"{target_name}{suffix}",
            )


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE_LIBRARY, OUT / "references_verified.csv")
    build_supplement()
    build_manuscript()
    build_cover_letter()
    copy_figures()
    build_repository()
    build_readme()
    build_revision_summary()
    print(f"Built IJMS package in {OUT}")
    print(f"Built GitHub package in {REPO}")
    print(f"Built Zenodo payload in {ZENODO}")


if __name__ == "__main__":
    main()
